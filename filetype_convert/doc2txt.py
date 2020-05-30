from docx import Document
import io
import shutil
import os

class FileTypeConvert(object):
    def __init__(self, file_name):
        self.file = file_name
        self.textFilename = ""

    def convertDocxToText(self):
        fileExtension = self.file.split(".")[-1]
        if fileExtension == "txt":
            self.textFilename = self.file
        elif fileExtension == "docx" or fileExtension == "doc":
            document = Document(self.file)
            self.textFilename = "".join(self.file.split(".")[:-1]) + ".txt"
            with io.open(self.textFilename, "w", encoding="utf-8") as textFile:
                for para in document.paragraphs:
                    x = str(para.text)
                    textFile.write((x)+'\n')

                # interate all the tables
                all_tables = document.tables
                for table in all_tables:
                    for row in table.rows:
                        for cell in row.cells:
                            textFile.write((x)+'\n')

                # iterate all the text box
                children = document.element.body.iter()
                for child in children:
                    if child.tag.endswith('textbox'):
                        for ci in child.iter():
                            if ci.tag.endswith('main}pPr'):
                                textFile.write('\n')
                            elif ci.tag.endswith('main}r'):
                                textFile.write(ci.text)

        else:
            print("Error: Not doc/docx type file!")
